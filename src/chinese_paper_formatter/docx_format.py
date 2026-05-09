"""DOCX rendering and reformatting for FormatSpec."""

from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from .format_spec import FormatSpec, HeadingSpec, ParagraphSpec
from .numbering import format_number
from .validation import validate_format_spec


def md_to_docx(input_path: Path, output_path: Path, spec: FormatSpec) -> Dict[str, Any]:
    """Render a lightly-structured Markdown manuscript to a new DOCX."""

    Document, _ = _docx_modules()
    blocks = parse_markdown(Path(input_path).read_text(encoding="utf-8"))
    doc = Document()
    report = _new_report(spec, input_path, output_path, mode="md-to-docx")
    _apply_page(doc, spec, report)
    counters = {1: 0, 2: 0, 3: 0}

    for block in blocks:
        kind = block["type"]
        if kind == "title":
            paragraph = doc.add_paragraph(block["text"])
            _apply_paragraph_style(paragraph, spec.front_matter.title, report)
            report["detected_blocks"].append("title")
        elif kind == "authors":
            paragraph = doc.add_paragraph(block["text"])
            _apply_paragraph_style(paragraph, spec.front_matter.authors, report)
            report["detected_blocks"].append("authors")
        elif kind == "abstract":
            paragraph = doc.add_paragraph(block["text"])
            _apply_paragraph_style(paragraph, spec.front_matter.abstract, report)
            report["detected_blocks"].append("abstract")
        elif kind == "keywords":
            paragraph = doc.add_paragraph(block["text"])
            _apply_paragraph_style(paragraph, spec.front_matter.keywords, report)
            report["detected_blocks"].append("keywords")
        elif kind == "heading":
            level = min(max(int(block["level"]), 1), 3)
            heading_spec = spec.headings.get(level)
            text = block["text"]
            if heading_spec and heading_spec.auto_number and not _has_heading_number(text):
                counters[level] += 1
                for reset_level in range(level + 1, 4):
                    counters[reset_level] = 0
                text = format_number(counters[level], heading_spec.numbering) + text
            paragraph = doc.add_paragraph(text)
            _apply_paragraph_style(paragraph, heading_spec or spec.paragraph, report)
            report["detected_blocks"].append(f"heading.{level}")
        elif kind == "reference_title":
            paragraph = doc.add_paragraph(block["text"])
            _apply_paragraph_style(paragraph, spec.references.title, report)
            report["detected_blocks"].append("references.title")
        elif kind == "reference_item":
            paragraph = doc.add_paragraph(_format_reference_text(block["text"], spec, report))
            _apply_reference_entry_style(paragraph, spec, report)
            report["detected_blocks"].append("references.entry")
        elif kind == "figure_caption":
            paragraph = doc.add_paragraph(block["text"])
            _apply_paragraph_style(paragraph, spec.figures_tables.figure_caption, report)
            report["detected_blocks"].append("figure_caption")
        elif kind == "table_caption":
            paragraph = doc.add_paragraph(block["text"])
            _apply_paragraph_style(paragraph, spec.figures_tables.table_caption, report)
            report["detected_blocks"].append("table_caption")
        elif kind == "table":
            _add_markdown_table(doc, block["rows"], report)
        else:
            paragraph = doc.add_paragraph(block["text"])
            _apply_paragraph_style(paragraph, spec.paragraph, report)
            report["detected_blocks"].append("body")

    _apply_advanced_word(spec, report)
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    report["applied"].append("saved output docx")
    return report


def apply_docx_format(input_path: Path, output_path: Path, spec: FormatSpec) -> Dict[str, Any]:
    """Apply a FormatSpec to an existing DOCX and save a new file."""

    Document, _ = _docx_modules()
    doc = Document(str(input_path))
    report = _new_report(spec, input_path, output_path, mode="docx-to-docx")
    _apply_page(doc, spec, report)
    in_references = False
    seen_first_text = False

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        kind, level = _classify_docx_paragraph(paragraph, text, seen_first_text, in_references)
        seen_first_text = True

        if kind == "title":
            _apply_paragraph_style(paragraph, spec.front_matter.title, report)
        elif kind == "abstract":
            _apply_paragraph_style(paragraph, spec.front_matter.abstract, report)
        elif kind == "keywords":
            _apply_paragraph_style(paragraph, spec.front_matter.keywords, report)
        elif kind == "reference_title":
            in_references = True
            _apply_paragraph_style(paragraph, spec.references.title, report)
        elif kind == "reference_item":
            _replace_paragraph_text(paragraph, _format_reference_text(text, spec, report))
            _apply_reference_entry_style(paragraph, spec, report)
        elif kind == "figure_caption":
            _apply_paragraph_style(paragraph, spec.figures_tables.figure_caption, report)
        elif kind == "table_caption":
            _apply_paragraph_style(paragraph, spec.figures_tables.table_caption, report)
        elif kind == "heading":
            _apply_paragraph_style(paragraph, spec.headings.get(level, spec.paragraph), report)
        else:
            _apply_paragraph_style(paragraph, spec.paragraph, report)
        report["detected_blocks"].append(kind if level is None else f"{kind}.{level}")

    _apply_tables(doc, spec, report)
    _apply_advanced_word(spec, report)
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    report["applied"].append("saved output docx")
    return report


def parse_markdown(text: str) -> List[Dict[str, Any]]:
    """Parse the light Markdown contract used by the skill."""

    blocks: List[Dict[str, Any]] = []
    in_references = False
    table_rows: List[List[str]] = []

    def flush_table() -> None:
        nonlocal table_rows
        if table_rows:
            blocks.append({"type": "table", "rows": table_rows})
            table_rows = []

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            flush_table()
            continue
        if line.startswith("|") and line.endswith("|"):
            row = [cell.strip() for cell in line.strip("|").split("|")]
            if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in row):
                table_rows.append(row)
            continue
        flush_table()

        heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading_match:
            hashes, value = heading_match.groups()
            if len(hashes) == 1 and not blocks:
                blocks.append({"type": "title", "text": value})
            else:
                level = max(1, min(len(hashes) - 1, 3))
                if _is_reference_title(value):
                    in_references = True
                    blocks.append({"type": "reference_title", "text": value})
                else:
                    blocks.append({"type": "heading", "level": level, "text": value})
            continue

        normalized = re.sub(r"^\*\*(.+?)\*\*", r"\1", line)
        if re.match(r"^摘\s*要[:：]", normalized):
            blocks.append({"type": "abstract", "text": normalized})
        elif re.match(r"^关键词[:：]", normalized):
            blocks.append({"type": "keywords", "text": normalized})
        elif _is_reference_title(normalized):
            in_references = True
            blocks.append({"type": "reference_title", "text": normalized})
        elif in_references:
            blocks.append({"type": "reference_item", "text": normalized})
        elif _is_figure_caption(normalized):
            blocks.append({"type": "figure_caption", "text": normalized})
        elif _is_table_caption(normalized):
            blocks.append({"type": "table_caption", "text": normalized})
        else:
            blocks.append({"type": "body", "text": normalized})
    flush_table()
    return blocks


def save_report(report: Dict[str, Any], path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")


def _docx_modules():
    try:
        from docx import Document
        from docx.shared import Cm, Pt
    except ImportError as exc:
        raise RuntimeError("缺少 python-docx；请先安装依赖: pip install -e .") from exc
    return Document, {"Cm": Cm, "Pt": Pt}


def _new_report(spec: FormatSpec, input_path: Path, output_path: Path, mode: str) -> Dict[str, Any]:
    validation = validate_format_spec(spec)
    return {
        "mode": mode,
        "template": spec.meta.name,
        "input_path": str(input_path),
        "output_path": str(output_path),
        "applied": [],
        "warnings": [issue.to_dict() for issue in validation.warnings],
        "unresolved_items": [issue.to_dict() for issue in validation.unresolved_items],
        "detected_blocks": [],
    }


def _apply_page(doc, spec: FormatSpec, report: Dict[str, Any]) -> None:
    _, shared = _docx_modules()
    Cm = shared["Cm"]
    for section in doc.sections:
        if spec.page.paper_size.upper() == "A4":
            section.page_width = Cm(21)
            section.page_height = Cm(29.7)
        section.top_margin = _to_length(spec.page.top_margin, default_font_size="12pt")
        section.bottom_margin = _to_length(spec.page.bottom_margin, default_font_size="12pt")
        section.left_margin = _to_length(spec.page.left_margin, default_font_size="12pt")
        section.right_margin = _to_length(spec.page.right_margin, default_font_size="12pt")
        if spec.page.header:
            section.header.paragraphs[0].text = spec.page.header
        if spec.page.footer:
            section.footer.paragraphs[0].text = spec.page.footer
    report["applied"].append("page margins and paper size")
    if spec.page.page_number:
        report["warnings"].append(
            {"severity": "warning", "path": "page.page_number", "message": "页码规则已记录；第一版不生成动态 Word 页码域"}
        )
    if spec.page.columns and spec.page.columns != 1:
        report["warnings"].append(
            {"severity": "warning", "path": "page.columns", "message": "分栏规则已记录；第一版不修改 Word 分栏 XML"}
        )


def _apply_paragraph_style(paragraph, spec: ParagraphSpec, report: Dict[str, Any]) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p_pr = paragraph._p.get_or_add_pPr()
    snap_to_grid = p_pr.find(qn("w:snapToGrid"))
    if snap_to_grid is None:
        snap_to_grid = OxmlElement("w:snapToGrid")
        p_pr.append(snap_to_grid)
    snap_to_grid.set(qn("w:val"), "0")

    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    paragraph.alignment = align_map.get(spec.align, WD_ALIGN_PARAGRAPH.JUSTIFY)
    fmt = paragraph.paragraph_format
    fmt.line_spacing = spec.line_height
    fmt.first_line_indent = _to_length(spec.first_line_indent, spec.font.size)
    fmt.left_indent = _to_length(spec.left_indent, spec.font.size)
    fmt.right_indent = _to_length(spec.right_indent, spec.font.size)
    fmt.space_before = _to_length(spec.space_before, spec.font.size, zero_when_none=True)
    fmt.space_after = _to_length(spec.space_after, spec.font.size, zero_when_none=True)

    if not paragraph.runs and paragraph.text:
        paragraph.add_run()
    for run in paragraph.runs:
        run.font.name = spec.font.family
        run._element.rPr.rFonts.set(qn("w:eastAsia"), spec.font.family)
        if spec.font.latin_family:
            run._element.rPr.rFonts.set(qn("w:ascii"), spec.font.latin_family)
            run._element.rPr.rFonts.set(qn("w:hAnsi"), spec.font.latin_family)
        run.font.size = _to_pt(spec.font.size)
        run.font.bold = spec.font.bold
        run.font.italic = spec.font.italic
    if spec.character_spacing:
        report["warnings"].append(
            {"severity": "warning", "path": "paragraph.character_spacing", "message": "字符间距规则已记录；第一版按段落报告，不直接写 OOXML"}
        )


def _apply_reference_entry_style(paragraph, spec: FormatSpec, report: Dict[str, Any]) -> None:
    _apply_paragraph_style(paragraph, spec.references.entry, report)
    indent = _to_length(spec.references.hanging_indent, spec.references.entry.font.size)
    if indent is not None and indent.pt != 0:
        paragraph.paragraph_format.left_indent = indent
        paragraph.paragraph_format.first_line_indent = -indent


def _apply_tables(doc, spec: FormatSpec, report: Dict[str, Any]) -> None:
    if not doc.tables:
        return
    for table in doc.tables:
        try:
            table.style = "Table Grid"
        except Exception:
            pass
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _apply_paragraph_style(paragraph, spec.paragraph, report)
    report["applied"].append("table text style")
    if spec.figures_tables.three_line_table:
        report["warnings"].append(
            {"severity": "warning", "path": "figures_tables.three_line_table", "message": "三线表规则已记录；第一版使用 Table Grid，不强改边框"}
        )


def _add_markdown_table(doc, rows: List[List[str]], report: Dict[str, Any]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=max(len(row) for row in rows))
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            table.cell(i, j).text = cell_text
    report["detected_blocks"].append("table")


def _apply_advanced_word(spec: FormatSpec, report: Dict[str, Any]) -> None:
    advanced = spec.advanced_word
    for field_name in [
        "east_asian_line_break",
        "latin_line_break",
        "use_document_grid",
        "character_spacing",
        "raw_ooxml_hints",
    ]:
        value = getattr(advanced, field_name)
        if value not in (None, "", {}, [], False):
            report["warnings"].append(
                {
                    "severity": "warning",
                    "path": f"advanced_word.{field_name}",
                    "message": "该高级 Word 规则已记录；第一版不保证完全应用",
                }
            )
    if advanced.widow_control is not None:
        report["applied"].append("recorded widow_control preference")
    if advanced.keep_with_next is not None:
        report["applied"].append("recorded keep_with_next preference")


def _classify_docx_paragraph(paragraph, text: str, seen_first_text: bool, in_references: bool) -> Tuple[str, Optional[int]]:
    style_name = (paragraph.style.name or "").lower()
    if not seen_first_text:
        return "title", None
    if _is_reference_title(text):
        return "reference_title", None
    if in_references:
        return "reference_item", None
    if re.match(r"^摘\s*要[:：]", text):
        return "abstract", None
    if re.match(r"^关键词[:：]", text):
        return "keywords", None
    if _is_figure_caption(text):
        return "figure_caption", None
    if _is_table_caption(text):
        return "table_caption", None
    heading_match = re.match(r"heading\s+([1-3])", style_name)
    if heading_match:
        return "heading", int(heading_match.group(1))
    if re.match(r"^[一二三四五六七八九十]+[、.．]\s*", text):
        return "heading", 1
    if re.match(r"^（[一二三四五六七八九十]+）", text):
        return "heading", 2
    if re.match(r"^\d+[.．]\s*", text):
        return "heading", 3
    return "body", None


def _is_reference_title(text: str) -> bool:
    return text.strip().replace("：", "").replace(":", "") in {"参考文献", "参考资料"}


def _is_figure_caption(text: str) -> bool:
    return bool(re.match(r"^图\s*\d+|^图[一二三四五六七八九十]+", text))


def _is_table_caption(text: str) -> bool:
    return bool(re.match(r"^表\s*\d+|^表[一二三四五六七八九十]+", text))


def _has_heading_number(text: str) -> bool:
    return bool(
        re.match(r"^[一二三四五六七八九十]+、", text)
        or re.match(r"^（[一二三四五六七八九十]+）", text)
        or re.match(r"^\d+[.．]", text)
    )


def _to_pt(value: Optional[str]):
    _, shared = _docx_modules()
    Pt = shared["Pt"]
    if value is None:
        return None
    if isinstance(value, (int, float)):
        return Pt(float(value))
    raw = str(value).strip()
    number = float(re.sub(r"[^0-9.\-]", "", raw) or 0)
    if raw.endswith("pt"):
        return Pt(number)
    if raw.endswith("em"):
        return Pt(number * 12)
    if raw.endswith("cm"):
        return Pt(number * 28.3464567)
    if raw.endswith("mm"):
        return Pt(number * 2.83464567)
    return Pt(number)


def _to_length(value: Optional[str], default_font_size: str, zero_when_none: bool = False):
    _, shared = _docx_modules()
    Cm = shared["Cm"]
    Pt = shared["Pt"]
    if value in (None, ""):
        return Pt(0) if zero_when_none else None
    if isinstance(value, (int, float)):
        return Pt(float(value))
    raw = str(value).strip()
    number = float(re.sub(r"[^0-9.\-]", "", raw) or 0)
    if raw.endswith("cm"):
        return Cm(number)
    if raw.endswith("mm"):
        return Cm(number / 10)
    if raw.endswith("pt"):
        return Pt(number)
    if raw.endswith("em"):
        font_size = float(re.sub(r"[^0-9.\-]", "", default_font_size) or 12)
        return Pt(number * font_size)
    return Pt(number)


def _format_reference_text(text: str, spec: FormatSpec, report: Dict[str, Any]) -> str:
    profile = (spec.references.punctuation or {}).get("format_profile")
    if profile not in {"chinese_author_year_no_type_mark", "chinese_social_science_author_year"}:
        return text
    converted = _format_chinese_author_year_reference(text)
    if converted != text:
        if "chinese author-year reference cleanup" not in report["applied"]:
            report["applied"].append("chinese author-year reference cleanup")
    return converted


def _format_chinese_author_year_reference(text: str) -> str:
    original = text.strip()
    if not original:
        return original
    text = re.sub(r"\s+", " ", original)
    match = re.match(r"^(?P<authors>.+?)\.\s*(?P<title>.+?)\[(?P<type>[A-Z/]+)\]\.\s*(?P<rest>.+?)\.?$", text)
    if not match:
        return original

    authors = _normalize_chinese_authors(match.group("authors"))
    title = match.group("title").strip()
    ref_type = match.group("type")
    rest = _normalize_reference_punctuation(match.group("rest").strip())

    if ref_type == "J":
        converted = _convert_journal_reference(authors, title, rest)
    elif ref_type == "M":
        converted = _convert_book_reference(authors, title, rest)
    elif ref_type == "N":
        converted = _convert_newspaper_reference(authors, title, rest)
    elif ref_type in {"EB/OL", "EB"}:
        converted = _convert_web_reference(authors, title, rest)
    else:
        converted = ""

    return converted or original


def _normalize_reference_punctuation(text: str) -> str:
    text = text.replace("：", ":")
    text = text.replace("，", ",")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def _normalize_chinese_authors(authors: str) -> str:
    if re.search(r"[\u4e00-\u9fff]", authors):
        authors = re.sub(r"\s*,\s*", "、", authors.strip())
    return authors.strip()


def _convert_journal_reference(authors: str, title: str, rest: str) -> str:
    match = re.match(
        r"^(?P<journal>.+?),\s*(?P<year>\d{4}[a-z]?),?\s*(?P<volume>\d+)?(?:\((?P<issue>\d+)\))?(?::(?P<pages>[^.]+))?$",
        rest,
    )
    if not match:
        return ""
    year = match.group("year")
    journal = match.group("journal").strip()
    issue = match.group("issue")
    suffix = f"第{issue}期" if issue else ""
    return f"{authors}，{year}，《{title}》，《{journal}》{suffix}。"


def _convert_book_reference(authors: str, title: str, rest: str) -> str:
    match = re.match(r"^(?P<place>.+?):\s*(?P<publisher>.+?),\s*(?P<year>\d{4}[a-z]?)$", rest)
    if not match:
        return ""
    return f"{authors}，{match.group('year')}，《{title}》，{match.group('place').strip()}：{match.group('publisher').strip()}。"


def _convert_newspaper_reference(authors: str, title: str, rest: str) -> str:
    match = re.match(r"^(?P<newspaper>.+?),\s*(?P<date>\d{4})-(?P<month>\d{1,2})-(?P<day>\d{1,2})(?:\([^)]*\))?$", rest)
    if not match:
        return ""
    date = f"{int(match.group('month'))}月{int(match.group('day'))}日"
    return f"{authors}，{match.group('date')}，《{title}》，《{match.group('newspaper').strip()}》{date}。"


def _convert_web_reference(authors: str, title: str, rest: str) -> str:
    match = re.match(
        r"^(?:(?P<source>.+?),\s*)?(?P<date>\d{4}[a-z]?-\d{1,2}-\d{1,2})\.\s*(?P<url>https?://\S+)$",
        rest,
    )
    if not match:
        match = re.match(
            r"^(?P<url>https?://\S+),\s*(?P<date>\d{4}[a-z]?-\d{1,2}-\d{1,2})(?:/\d{4}-\d{1,2}-\d{1,2})?$",
            rest,
        )
    if not match:
        return ""
    year = match.group("date")[:4]
    source = match.groupdict().get("source")
    source_part = f"{source.strip()}，" if source else ""
    return f"{authors}，{year}，《{title}》，{source_part}{match.group('url')}，{_format_iso_date(match.group('date'))}。"


def _format_iso_date(value: str) -> str:
    year, month, day = re.match(r"^(\d{4}[a-z]?)-(\d{1,2})-(\d{1,2})", value).groups()
    return f"{year}年{int(month)}月{int(day)}日"


def _replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.text == text:
        return
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)
